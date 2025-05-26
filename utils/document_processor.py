import os
from typing import List, Dict, Any, Tuple
from pathlib import Path
import PyPDF2
from docx import Document
from PIL import Image
import base64
import io

class DocumentProcessor:
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_docx,
            '.txt': self._process_txt,
            '.png': self._process_image,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
            '.gif': self._process_image
        }
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Process any supported document type"""
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        processor = self.supported_formats[file_extension]
        return processor(file_path)
    
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text and metadata from PDF"""
        result = {
            'type': 'pdf',
            'text_content': [],
            'images': [],
            'metadata': {}
        }
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                result['metadata'] = {
                    'pages': len(pdf_reader.pages),
                    'title': pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                    'author': pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else ''
                }
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        result['text_content'].append({
                            'page': page_num + 1,
                            'content': text.strip(),
                            'char_count': len(text)
                        })
                    
                    # Check for images (basic detection)
                    if '/XObject' in page.get('/Resources', {}):
                        result['images'].append({
                            'page': page_num + 1,
                            'detected': True,
                            'extractable': False  # PDF image extraction is complex
                        })
        
        except Exception as e:
            result['error'] = str(e)
        
        return result
    
    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text and images from DOCX"""
        result = {
            'type': 'docx',
            'text_content': [],
            'images': [],
            'metadata': {}
        }
        
        try:
            doc = Document(file_path)
            
            # Extract metadata
            result['metadata'] = {
                'paragraphs': len(doc.paragraphs),
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or ''
            }
            
            # Extract text content
            for para_num, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    result['text_content'].append({
                        'paragraph': para_num + 1,
                        'content': paragraph.text.strip(),
                        'char_count': len(paragraph.text)
                    })
            
            # Check for inline shapes (images)
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    result['images'].append({
                        'detected': True,
                        'type': 'inline_image',
                        'rel_id': rel.rId
                    })
        
        except Exception as e:
            result['error'] = str(e)
        
        return result
    
    def _process_txt(self, file_path: str) -> Dict[str, Any]:
        """Process plain text files"""
        result = {
            'type': 'txt',
            'text_content': [],
            'images': [],
            'metadata': {}
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                result['text_content'].append({
                    'content': content,
                    'char_count': len(content)
                })
                result['metadata'] = {
                    'size': len(content),
                    'lines': content.count('\n') + 1
                }
        except Exception as e:
            result['error'] = str(e)
        
        return result
    
    def _process_image(self, file_path: str) -> Dict[str, Any]:
        """Process image files"""
        result = {
            'type': 'image',
            'text_content': [],
            'images': [],
            'metadata': {}
        }
        
        try:
            with Image.open(file_path) as img:
                # Convert to base64 for API calls
                buffer = io.BytesIO()
                img.save(buffer, format=img.format or 'PNG')
                img_str = base64.b64encode(buffer.getvalue()).decode()
                
                result['images'].append({
                    'base64': img_str,
                    'format': img.format,
                    'size': img.size,
                    'mode': img.mode
                })
                
                result['metadata'] = {
                    'width': img.width,
                    'height': img.height,
                    'format': img.format,
                    'mode': img.mode
                }
        except Exception as e:
            result['error'] = str(e)
        
        return result
